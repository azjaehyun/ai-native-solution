import boto3
import json
import faiss
import numpy as np
import base64
from typing import List
from io import BytesIO
import sys


print(sys.path)


# (1) Word íŒŒì¼ ì²˜ë¦¬ìš© ë¼ì´ë¸ŒëŸ¬ë¦¬
try:
    from docx import Document
    print("docx ë¼ì´ë¸ŒëŸ¬ë¦¬ import ì„±ê³µ!")
except ImportError as e:
    print(f"docx ImportError ë°œìƒ: {e}")

# (2) PDF íŒŒì¼ ì²˜ë¦¬ìš© ë¼ì´ë¸ŒëŸ¬ë¦¬
try:
    import PyPDF2
    print("PyPDF2 ë¼ì´ë¸ŒëŸ¬ë¦¬ import ì„±ê³µ!")
except ImportError:
    pass

# (3) Excel íŒŒì¼ ì²˜ë¦¬ìš© ë¼ì´ë¸ŒëŸ¬ë¦¬
try:
    from openpyxl import load_workbook
    print("openpyxl ë¼ì´ë¸ŒëŸ¬ë¦¬ import ì„±ê³µ!")
except ImportError:
    pass

# Bedrock í´ë¼ì´ì–¸íŠ¸ ìƒì„±
bedrock_agent_runtime = boto3.client(
    service_name="bedrock-agent-runtime",
    region_name="ap-northeast-2"  # ì‹¤ì œ ì‚¬ìš©í•˜ëŠ” ë¦¬ì „ìœ¼ë¡œ ë³€ê²½
)

def dynamic_chunk_text(
    text: str,
    strategy: str = "fixed",
    chunk_size: int = 500,
    overlap_size: int = 50
) -> List[str]:
    text = text.strip()
    if not text:
        return []
        

    if strategy == "hybrid":
        # Combine benefits of all strategies
        import re
        chunks = []

        # Split by paragraphs
        paragraphs = text.split('\n\n')
        for para in paragraphs:
            para = para.strip()
            if len(para) <= chunk_size:
                if para:
                    chunks.append(para)
            else:
                # Further split large paragraphs by sentences
                sentences = re.split(r'(?<=[.!?])\s+', para)
                for sent in sentences:
                    sent = sent.strip()
                    if len(sent) <= chunk_size:
                        if sent:
                            chunks.append(sent)
                    else:
                        # Use sliding window for long sentences
                        start = 0
                        while start < len(sent):
                            end = start + chunk_size
                            chunk = sent[start:end]
                            chunks.append(chunk)
                            start += max(chunk_size - overlap_size, 1)

        return chunks

    elif strategy == "paragraph":
        paragraphs = text.split('\n\n')
        chunks = []
        for para in paragraphs:
            para = para.strip()
            if len(para) <= chunk_size:
                if para:
                    chunks.append(para)
            else:
                start = 0
                while start < len(para):
                    end = start + chunk_size
                    chunk = para[start:end]
                    chunks.append(chunk)
                    start += max(chunk_size - overlap_size, 1)
        return chunks

    elif strategy == "sentence":
        import re
        sentences = re.split(r'(?<=[.!?])\s+', text)
        chunks = []
        for sent in sentences:
            sent = sent.strip()
            if len(sent) <= chunk_size:
                if sent:
                    chunks.append(sent)
            else:
                start = 0
                while start < len(sent):
                    end = start + chunk_size
                    chunk = sent[start:end]
                    chunks.append(chunk)
                    start += max(chunk_size - overlap_size, 1)
        return chunks

    elif strategy == "token":
        words = text.split()
        chunks = []
        start_idx = 0
        while start_idx < len(words):
            end_idx = start_idx + chunk_size
            chunk_words = words[start_idx:end_idx]
            chunk = " ".join(chunk_words)
            chunks.append(chunk)
            start_idx += max(chunk_size - overlap_size, 1)
        return chunks

    elif strategy == "sliding_window":
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += max(chunk_size - overlap_size, 1)
        return chunks

    elif strategy == "by_section":
        sections = text.split("\n\n\n")  # Split by large section gaps
        chunks = []
        for section in sections:
            section = section.strip()
            if len(section) <= chunk_size:
                if section:
                    chunks.append(section)
            else:
                start = 0
                while start < len(section):
                    end = start + chunk_size
                    chunk = section[start:end]
                    chunks.append(chunk)
                    start += max(chunk_size - overlap_size, 1)
        return chunks

    elif strategy == "by_keyword":
        keywords = ["Introduction", "Summary", "Conclusion"]  # Example keywords
        chunks = []
        for keyword in keywords:
            start_idx = text.find(keyword)
            while start_idx != -1:
                end_idx = start_idx + chunk_size
                chunk = text[start_idx:end_idx]
                chunks.append(chunk)
                start_idx = text.find(keyword, start_idx + 1)
        return chunks

    else:
        chunks = []
        start = 0
        while start < len(text):
            end = start + chunk_size
            chunk = text[start:end]
            chunks.append(chunk)
            start += chunk_size
        return chunks


def chunk_word(
    file_bytes: bytes,
    chunk_size: int = 500,
    overlap_size: int = 50
) -> List[str]:
    chunks = []
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        print(f"Error loading Word docx: {e}")
        return []

    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if not text:
            continue

        if len(text) <= chunk_size:
            chunks.append(f"[Paragraph {i+1}] {text}")
        else:
            start = 0
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                chunks.append(f"[Paragraph {i+1}] {chunk}")
                start += max(chunk_size - overlap_size, 1)

    return chunks


def chunk_word_heading(
    file_bytes: bytes,
    chunk_size: int = 500,
    overlap_size: int = 50
) -> List[str]:
    chunks = []
    try:
        doc = Document(BytesIO(file_bytes))
    except Exception as e:
        print(f"Error loading Word docx: {e}")
        return []

    all_paragraphs = doc.paragraphs
    current_heading = None
    current_text_buffer = []

    def flush_chunk(heading, text_buffer):
        if not text_buffer:
            return []
        combined_text = f"[Heading: {heading}] " + "\n".join(text_buffer)
        if len(combined_text) <= chunk_size:
            return [combined_text]
        else:
            splitted_chunks = []
            start = 0
            while start < len(combined_text):
                end = start + chunk_size
                chunk_part = combined_text[start:end]
                splitted_chunks.append(chunk_part)
                start += max(chunk_size - overlap_size, 1)
            return splitted_chunks

    for para in all_paragraphs:
        style_name = getattr(para.style, 'name', '') or ''
        text = para.text.strip()
        if 'Heading' in style_name:
            if current_heading or current_text_buffer:
                chunks.extend(flush_chunk(current_heading, current_text_buffer))
            current_heading = text
            current_text_buffer = []
        else:
            if not current_heading:
                current_heading = "No Heading"
            if text:
                current_text_buffer.append(text)

    if current_heading or current_text_buffer:
        chunks.extend(flush_chunk(current_heading, current_text_buffer))

    return chunks


def chunk_pdf(
    file_bytes: bytes,
    chunk_size: int = 500,
    overlap_size: int = 50
) -> List[str]:
    chunks = []
    try:
        pdf_reader = PyPDF2.PdfReader(BytesIO(file_bytes))
    except Exception as e:
        print(f"Error loading PDF: {e}")
        return []

    for page_index, page in enumerate(pdf_reader.pages):
        try:
            text = page.extract_text()
        except:
            text = ""
        text = text.strip() if text else ""
        if not text:
            continue

        if len(text) <= chunk_size:
            chunks.append(f"[Page {page_index+1}] {text}")
        else:
            start = 0
            while start < len(text):
                end = start + chunk_size
                chunk = text[start:end]
                chunks.append(f"[Page {page_index+1}] {chunk}")
                start += max(chunk_size - overlap_size, 1)

    return chunks

def chunk_excel(
    file_bytes: bytes,
    chunk_size: int = 500,
    overlap_size: int = 50,
    header_row_index: int = 1
) -> list:
    """
    Excel íŒŒì¼ì„ ì‹œíŠ¸(íƒ­) â†’ í–‰(Row) ìˆœìœ¼ë¡œ ìˆœíšŒ í›„,
    ì—¬ëŸ¬ í–‰ì„ í•©ì³ í•˜ë‚˜ì˜ ì²­í¬ë¥¼ ìƒì„±í•˜ë©°,
    ê° ì²­í¬ì˜ ì²« ë¶€ë¶„ì— ì§€ì •ëœ í—¤ë” í–‰ì„ ë°˜ë³µì ìœ¼ë¡œ í¬í•¨.

    Args:
        file_bytes (bytes): Excel íŒŒì¼ ë°”ì´íŠ¸ ë°ì´í„°
        chunk_size (int): ìµœëŒ€ ì²­í¬ í¬ê¸° (ê¸°ë³¸ê°’ 500)
        overlap_size (int): ì²­í¬ ê°„ ì¤‘ì²© í¬ê¸° (ê¸°ë³¸ê°’ 50)
        header_row_index (int): í—¤ë” í–‰ì˜ ì¸ë±ìŠ¤ (ê¸°ë³¸ê°’ 1, ì²« ë²ˆì§¸ í–‰)

    Returns:
        List[str]: ì²­í¬ ë¦¬ìŠ¤íŠ¸
    """
    chunks = []
    try:
        wb = load_workbook(filename=BytesIO(file_bytes), read_only=False, data_only=False)
    except Exception as e:
        print(f"Error loading Excel: {e}")
        return []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        print(f"Processing sheet: {sheet_name}")

        header = None  # í—¤ë” ê°’ì„ ì €ì¥
        current_chunk = ""

        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            # ì§€ì •ëœ í—¤ë” í–‰ì„ ì²˜ë¦¬
            if row_idx == header_row_index:
                header = " | ".join([str(cell) if cell is not None else '' for cell in row])
                continue

            # í—¤ë” ì´í›„ì˜ ë°ì´í„°ë¥¼ ì²˜ë¦¬
            if row_idx > header_row_index:
                # í–‰ ë°ì´í„° í•©ì¹˜ê¸°
                row_str_list = [str(cell) if cell is not None else '' for cell in row]
                row_text = " | ".join(row_str_list).strip()

                if not row_text:
                    continue

                # í˜„ì¬ ì²­í¬ì— í–‰ ë°ì´í„° ì¶”ê°€
                if len(current_chunk) + len(row_text) + 1 <= chunk_size:
                    current_chunk += f"{row_text}\n"
                else:
                    # í˜„ì¬ ì²­í¬ë¥¼ ì €ì¥í•˜ê³  ìƒˆ ì²­í¬ ì‹œì‘
                    chunks.append(f"[Sheet: {sheet_name}]\n{header}\n{current_chunk.strip()}")
                    # ì¤‘ì²©ëœ ë°ì´í„°ë¥¼ í¬í•¨í•œ ìƒˆë¡œìš´ ì²­í¬ ìƒì„±
                    if overlap_size > 0 and len(current_chunk) > overlap_size:
                        current_chunk = current_chunk[-overlap_size:] + row_text + "\n"
                    else:
                        current_chunk = row_text + "\n"

        # ë‚¨ì•„ìˆëŠ” ë°ì´í„° ì²˜ë¦¬
        if current_chunk.strip():
            chunks.append(f"[Sheet: {sheet_name}]\n{header}\n{current_chunk.strip()}")

    wb.close()
    return chunks



def chunk_excel_domain_based(
    file_bytes: bytes,
    category_col: int = 0
) -> List[str]:
    chunks = []
    try:
        # domain-basedì—ì„œëŠ” ì—¬ì „íˆ í•„ìš”ì— ë”°ë¼ data_only=False ì‚¬ìš© ê°€ëŠ¥
        wb = load_workbook(filename=BytesIO(file_bytes), read_only=False, data_only=False)
    except Exception as e:
        print(f"Error loading Excel for domain chunking: {e}")
        return []

    for sheet_name in wb.sheetnames:
        sheet = wb[sheet_name]
        print(f"[Domain-based] Processing sheet: {sheet_name}")
        
        category_map = {}

        for row_idx, row in enumerate(sheet.iter_rows(values_only=True), start=1):
            # (ì„ íƒ) í—¤ë” ìŠ¤í‚µ
            # if row_idx == 1:
            #     continue

            if not row or all(cell is None for cell in row):
                continue

            cat_val = row[category_col] if category_col < len(row) else None
            cat_val = str(cat_val) if cat_val is not None else "UnknownCategory"

            row_str_list = [str(cell) if cell is not None else '' for cell in row]
            row_text = " | ".join(row_str_list).strip()

            if cat_val not in category_map:
                category_map[cat_val] = []
            category_map[cat_val].append(row_text)

        for cat_val, rows in category_map.items():
            chunk_text = f"[Sheet: {sheet_name} / Category: {cat_val}]\n" + "\n".join(rows)
            chunks.append(chunk_text)

    wb.close()
    return chunks


def detect_file_type(filename: str) -> str:
    lower_name = filename.lower()
    if lower_name.endswith('.docx'):
        return 'word'
    elif lower_name.endswith('.xlsx') or lower_name.endswith('.xls'):
        return 'excel'
    elif lower_name.endswith('.pdf'):
        return 'pdf'
    else:
        return 'text'


def chunk_word_full_text(
    file_bytes: bytes,
    chunk_size: int = 500,
    overlap_size: int = 50
) -> List[str]:
    try:
        doc = Document(BytesIO(file_bytes))
        print("[INFO] Word document successfully loaded.")
    except Exception as e:
        print(f"[ERROR] Error loading Word docx: {e}")
        return []

    all_text = []

    # Collect text from paragraphs
    print("[INFO] Extracting paragraphs from the document...")
    for i, para in enumerate(doc.paragraphs):
        text = para.text.strip()
        if text:
            all_text.append(text)
            print(f"[DEBUG] Paragraph {i + 1}: {text}")

    # Include table content if any
    print("[INFO] Extracting tables from the document...")
    for table_idx, table in enumerate(doc.tables):
        print(f"[DEBUG] Table {table_idx + 1} found.")
        for row_idx, row in enumerate(table.rows):
            row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_text:
                table_row_text = " | ".join(row_text)
                all_text.append(table_row_text)
                print(f"[DEBUG] Table {table_idx + 1}, Row {row_idx + 1}: {table_row_text}")

    # Collect text from headers and footers if any
    print("[INFO] Extracting headers and footers from the document...")
    if doc.sections:
        for section_idx, section in enumerate(doc.sections):
            print(f"[DEBUG] Section {section_idx + 1} found.")
            header = section.header
            footer = section.footer
            if header:
                for para_idx, para in enumerate(header.paragraphs):
                    text = para.text.strip()
                    if text:
                        header_text = f"[Header] {text}"
                        all_text.append(header_text)
                        print(f"[DEBUG] Section {section_idx + 1} Header Paragraph {para_idx + 1}: {header_text}")
            if footer:
                for para_idx, para in enumerate(footer.paragraphs):
                    text = para.text.strip()
                    if text:
                        footer_text = f"[Footer] {text}"
                        all_text.append(footer_text)
                        print(f"[DEBUG] Section {section_idx + 1} Footer Paragraph {para_idx + 1}: {footer_text}")

    # Combine all extracted text
    print("[INFO] Combining all extracted text...")
    combined_text = "\n".join(all_text)
    print(f"[DEBUG] Combined Text Length: {len(combined_text)} characters")

    # Split into chunks
    print("[INFO] Splitting text into chunks...")
    chunks = []
    start = 0
    while start < len(combined_text):
        end = start + chunk_size
        chunk = combined_text[start:end]
        chunks.append(chunk)
        print(f"[DEBUG] Chunk {len(chunks)}: {chunk[:50]}...")  # Print first 50 chars of each chunk
        start += max(chunk_size - overlap_size, 1)

    print(f"[INFO] Total Chunks Created: {len(chunks)}")
    return chunks

def read_and_chunk_file(
    file_bytes: bytes,
    filename: str,
    chunk_strategy: str = None,
    chunk_size: int = 500,
    overlap_size: int = 50,
    category_col: int = 0
) -> List[str]:
    print(f"[DEBUG] Starting read_and_chunk_file with filename={filename}, chunk_strategy={chunk_strategy}, chunk_size={chunk_size}, overlap_size={overlap_size}, category_col={category_col}")
    file_type = detect_file_type(filename)
    print(f"[DEBUG] Detected file type: {file_type}")

    if file_type == 'word':
        # Default strategy for Word files
        if chunk_strategy is None:
            chunk_strategy = "full_text"
        print(f"[DEBUG] Using chunk_strategy={chunk_strategy} for Word file")

        if chunk_strategy == "heading":
            return chunk_word_heading(file_bytes, chunk_size, overlap_size)
        elif chunk_strategy == "full_text":
            return chunk_word_full_text(file_bytes, chunk_size, overlap_size)
        else:
            return chunk_word(file_bytes, chunk_size, overlap_size)

    elif file_type == 'excel':
        # Default strategy for Excel files
        if chunk_strategy is None:
            chunk_strategy = "chunk_excel"
        print(f"[DEBUG] Using chunk_strategy={chunk_strategy} for Excel file")

        if chunk_strategy == "chunk_excel":
            return chunk_excel(file_bytes, chunk_size, overlap_size)
        else:
            return chunk_excel_domain_based(file_bytes, category_col=category_col)

    elif file_type == 'pdf':
        # Default strategy for PDF files
        if chunk_strategy is None:
            chunk_strategy = "hybrid"
        print(f"[DEBUG] Using chunk_strategy={chunk_strategy} for PDF file")

        if chunk_strategy == "hybrid":
            # Combine PDF chunking with dynamic chunking
            text_chunks = chunk_pdf(file_bytes, chunk_size, overlap_size)
            print(f"[DEBUG] PDF chunking produced {len(text_chunks)} chunks.")
            combined_text = "\n".join(text_chunks)
            return dynamic_chunk_text(combined_text, strategy=chunk_strategy, chunk_size=chunk_size, overlap_size=overlap_size)
        else:
            return chunk_pdf(file_bytes, chunk_size, overlap_size)

    else:
        # Default strategy for plain text files
        if chunk_strategy is None:
            chunk_strategy = "hybrid"
        print(f"[DEBUG] Using chunk_strategy={chunk_strategy} for plain text file")

        decoded_text = file_bytes.decode('utf-8', errors='ignore')
        return dynamic_chunk_text(decoded_text, strategy=chunk_strategy, chunk_size=chunk_size, overlap_size=overlap_size)
    

def generate_embeddings(text: List[str], model_id: str) -> List[np.ndarray]:
    print(f"Generating embeddings for {len(text)} chunks using model {model_id}...")
    bedrock_runtime = boto3.client('bedrock-runtime', region_name='ap-northeast-2')
    embeddings = []

    for i, chunk in enumerate(text):
        try:
            #print(f"[{i + 1}/{len(text)}] Processing chunk: {chunk[:50]}...")
            print(f"[{i + 1}/{len(text)}] Processing chunk: {chunk} ")
            cleaned_chunk = chunk.replace('"', '\\"').strip()
            if not cleaned_chunk:
                print(f"Skipping empty or invalid chunk at index {i}.")
                continue

            payload = {
                "inputText": cleaned_chunk
            }

            response = bedrock_runtime.invoke_model(
                modelId=model_id,
                contentType='application/json',
                accept='application/json',
                body=json.dumps(payload)
            )

            response_body = json.loads(response['body'].read())
            if 'embedding' not in response_body:
                raise ValueError("Response does not contain 'embedding' field.")

            embedding = np.array(response_body['embedding'], dtype=np.float32)
            embeddings.append(embedding)
            print(f"Chunk {i + 1} processed successfully.")

        except Exception as e:
            print(f"Error processing chunk {i + 1}: {e}")
            continue

    print("Embeddings generation completed.")
    return embeddings


def create_faiss_index(embeddings: List[np.ndarray]) -> faiss.IndexFlatL2:
    print("Creating FAISS index...")
    dimension = embeddings[0].shape[0]
    index = faiss.IndexFlatL2(dimension)
    index.add(np.vstack(embeddings))
    print("FAISS index created.")
    return index

def print_retrieval_details(indices: np.ndarray, distances: np.ndarray, texts: List[str]) -> None:
    """
    Print detailed retrieval results including indices, distances, and corresponding texts.

    Args:
        indices (np.ndarray): Indices of the retrieved items.
        distances (np.ndarray): Distances of the retrieved items.
        texts (List[str]): List of texts from which items are retrieved.
    """
    print("Detailed retrieval results:")
    for idx, dist in zip(indices[0], distances[0]):
        if idx < len(texts):
            print(f"Index: {idx}, Distance: {dist:.4f}, Text: {texts[idx]}")
        else:
            print(f"Index: {idx}, Distance: {dist:.4f}, Text: [Index out of bounds]")


def retrieve_relevant_chunks(query: str, faiss_index: faiss.IndexFlatL2, texts: List[str], model_id: str, top_k: int = 3) -> List[str]:
    print(f"Retrieving relevant chunks for query: {query}")
    query_embedding = generate_embeddings([query], model_id=model_id)[0]
    distances, indices = faiss_index.search(np.array([query_embedding], dtype=np.float32), top_k)
    print(f"Retrieved indices: {indices[0]}, distances: {distances[0]}")
    print_retrieval_details(indices, distances, texts)
    return [texts[i] for i in indices[0] if i < len(texts)]


def lambda_handler(event, context):
    print("Lambda function invoked...")
    try:
        body = json.loads(event.get('body', '{}'))
        print("Request body:", body)

        file_content_base64 = body.get('fileContent', None)
        filename = body.get('fileName', 'default.md')
        new_message = body.get('message', 'ìŠ¤ë…¸ìš°í”Œë ˆì´í¬ ë¬¸ì œ í•œê°œë§Œ ê°ê´€ì‹ìœ¼ë¡œ ë§Œë“¤ì–´ì¤˜')
        history = body.get('chatHistoryMessage', [])
        selectedModel = body.get('selectedModel', 'claude3.5')

        chunk_strategy = body.get('chunkStrategy', None)  #  fixed
        max_chunk_size = body.get('maxChunkSize', 500)
        if 'overlapSize' not in body or body['overlapSize'] is None:
            overlap_size = int(max_chunk_size * 0.1)
            print(f"overlapSize not provided. Using 10% of maxChunkSize => {overlap_size}")
        else:
            overlap_size = body['overlapSize']

        category_col = body.get('categoryCol', 0)

        print(f"chunkStrategy={chunk_strategy}, maxChunkSize={max_chunk_size}, overlapSize={overlap_size}, categoryCol={category_col}")

        embedding_model_id = 'amazon.titan-embed-text-v2:0'
        kbId = 'ZXCWNTBUPU'

        if selectedModel == 'claude3.5':
            modelArn = 'arn:aws:bedrock:ap-northeast-2::foundation-model/anthropic.claude-3-5-sonnet-20240620-v1:0'
        elif selectedModel == 'claude3.0':
            modelArn = 'arn:aws:bedrock:ap-northeast-2::foundation-model/anthropic.claude-3-sonnet-20240229-v1:0'
        else:
            raise ValueError(f"Unsupported model selected: {selectedModel}")

        print(f"Selected model ARN: {modelArn}")
        relevant_chunks = []

        if file_content_base64:
            try:
                file_content = base64.b64decode(file_content_base64)
            except Exception as e:
                raise ValueError(f"Error decoding base64 file content: {str(e)}")

            print("Processing uploaded file...")
            text_chunks = read_and_chunk_file(
                file_bytes=file_content,
                filename=filename,
                chunk_strategy=chunk_strategy,
                chunk_size=max_chunk_size,
                overlap_size=overlap_size,
                category_col=category_col
            )

            print(f"Total chunks created: {len(text_chunks)}")

            if text_chunks:
                embeddings = generate_embeddings(text_chunks, embedding_model_id)
                faiss_index = create_faiss_index(embeddings)
                relevant_chunks = retrieve_relevant_chunks(new_message, faiss_index, text_chunks, embedding_model_id)
        else:
            print("No fileContent provided. Skipping file chunking and retrieval steps.")

        prompt_template = {
            "system_prompt": (
                "YOU ARE A HIGHLY ADVANCED LANGUAGE MODEL DEPLOYED ON AWS BEDROCK, "
                "DESIGNED TO PROVIDE EXPERT-LEVEL ASSISTANCE IN A WIDE VARIETY OF TASKS. "
                "YOUR OBJECTIVE IS TO DELIVER RELIABLE, ACCURATE, AND CONTEXTUALLY APPROPRIATE RESPONSES BASED ON THE USER'S INPUT. "
                "YOU MUST MAINTAIN PROFESSIONALISM AND ADAPTABILITY IN ALL INTERACTIONS.\n\n"
                "### INSTRUCTIONS ###\n"
                "1. **COMPREHEND THE USER'S REQUEST**: Carefully analyze the input to understand the user's intent and context.\n"
                "2. **PROVIDE STRUCTURED AND ACCURATE RESPONSES**:\n"
                "   - Break down complex tasks into manageable steps.\n"
                "   - Use logical reasoning and domain-specific knowledge to formulate answers.\n"
                "   - Adapt your response based on the user's expertise and context.\n"
                "3. **OPTIMIZE FOR CLARITY AND RELEVANCE**:\n"
                "   - Use simple language for general users or technical terms for experts as appropriate.\n"
                "   - Incorporate examples, analogies, or step-by-step guides to enhance understanding.\n"
                "4. **ENSURE FLEXIBILITY AND CREATIVITY**:\n"
                "   - Address requests that require innovation or unique solutions.\n"
                "   - Provide options when suitable to accommodate varying user needs.\n"
                "5. **HANDLE EDGE CASES AND LIMITATIONS**:\n"
                "   - Identify potential exceptions and clarify boundaries of the provided solution.\n"
                "   - If uncertain, explicitly state the limitation or offer alternative approaches.\n\n"
                "### FOR CASUAL AND EVERYDAY CONVERSATIONS ###\n"
                "1. **USE A FRIENDLY AND WARM TONE**: Respond in a conversational and approachable manner.\n"
                "2. **DEMONSTRATE EMPATHY AND UNDERSTANDING**: Reflect emotional intelligence by acknowledging the user's feelings and context.\n"
                "3. **ADAPT TO INFORMAL CONTEXTS**: Use simpler and more natural language for casual chats (e.g., 'ì•ˆë…•!' or 'ì–´ë–»ê²Œ ì§€ë‚´ì„¸ìš”?').\n"
                "4. **INCLUDE POSITIVE REINFORCEMENT**: Add uplifting or encouraging remarks when appropriate (e.g., 'ì¢‹ì€ í•˜ë£¨ ë³´ë‚´ì„¸ìš”!' or 'ì •ë§ ë©‹ì§„ ìƒê°ì´ì—ìš”!').\n"
                "5. **BE ENGAGING AND RESPONSIVE**: Ask follow-up questions to keep the conversation lively and engaging.\n"
                "6. **STAY NEUTRAL AND SUPPORTIVE**: Ensure responses remain polite, neutral, and culturally sensitive in all situations.\n\n"
                "### CHAIN OF THOUGHT ###\n"
                "1. **UNDERSTAND**: Carefully interpret the user's input, clarifying ambiguities if necessary.\n"
                "2. **BASICS**: Identify the fundamental concepts or components involved in the request.\n"
                "3. **BREAK DOWN**: Divide the task into smaller, logical parts for deeper analysis.\n"
                "4. **ANALYZE**: Apply domain-specific knowledge and logic to evaluate each part.\n"
                "5. **BUILD**: Synthesize the findings into a comprehensive and coherent response.\n"
                "6. **EDGE CASES**: Consider unusual or extreme scenarios to ensure robustness.\n"
                "7. **FINAL OUTPUT**: Present the solution clearly, with actionable steps if applicable.\n\n"
                "### WHAT NOT TO DO ###\n"
                "- DO NOT PROVIDE INACCURATE OR UNSUPPORTED INFORMATION.\n"
                "- DO NOT OVERCOMPLICATE ANSWERS UNNECESSARILY.\n"
                "- DO NOT ASSUME USER INTENT WITHOUT SUFFICIENT CONTEXT.\n"
                "- DO NOT OMIT IMPORTANT DETAILS OR FAIL TO ADDRESS ALL ASPECTS OF A REQUEST.\n"
                "- DO NOT GENERATE OFFENSIVE, BIASED, OR INAPPROPRIATE CONTENT.\n\n"
                "### OUTPUT STYLE GUIDELINES ###\n"
                "- MAINTAIN A PROFESSIONAL AND NEUTRAL TONE AT ALL TIMES.\n"
                "- FORMAT RESPONSES CLEARLY USING HEADINGS, BULLET POINTS, OR NUMBERED LISTS WHEN APPROPRIATE.\n"
                "- INCLUDE EXAMPLES, ANALOGIES, OR VISUALIZATION IDEAS TO IMPROVE COMPREHENSION.\n"
                "- FOR CASUAL INTERACTIONS, USE A LIGHTER, MORE RELAXED TONE TO MAKE THE USER FEEL AT EASE.\n\n"
                "### EXAMPLES OF USAGE ###\n"
                "1. **GENERAL INQUIRY**:\n"
                "   - Input: \"Explain the basics of machine learning.\"\n"
                "   - Output: \"Machine learning is a subset of AI where algorithms learn patterns from data... (detailed explanation). Example: A spam filter learns to classify emails...\"\n\n"
                "2. **CASUAL GREETING**:\n"
                "   - Input: \"ì•ˆë…•?\"\n"
                "   - Output: \"ì•ˆë…•í•˜ì„¸ìš”! ì˜¤ëŠ˜ í•˜ë£¨ëŠ” ì–´ë• ì–´ìš”? ğŸ˜Š\"\n\n"
                "3. **EMPATHETIC RESPONSE**:\n"
                "   - Input: \"ê¸°ë¶„ì´ ì¢€ ì•ˆ ì¢‹ì•„.\"\n"
                "   - Output: \"ë¬´ìŠ¨ ì¼ì´ ìˆìœ¼ì‹ ê°€ìš”? ê´œì°®ìœ¼ì‹œë‹¤ë©´ ì´ì•¼ê¸° ë‚˜ëˆ ë³´ì„¸ìš”. ì œê°€ ë„ìš¸ ìˆ˜ ìˆëŠ” ì¼ì´ ìˆì„ì§€ë„ ëª°ë¼ìš”. ğŸ’›\"\n\n"
                "4. **CREATIVE TASKS**:\n"
                "   - Input: \"Suggest three ideas for a marketing campaign.\"\n"
                "   - Output: \"1. A social media challenge that involves... 2. Influencer partnerships focused on... 3. Interactive content like quizzes...\"\n\n"
                "### ADAPTABILITY ###\n"
                "THIS PROMPT SHOULD ENABLE YOU TO HANDLE:\n"
                "- GENERAL KNOWLEDGE QUESTIONS\n"
                "- TECHNICAL SUPPORT AND GUIDANCE\n"
                "- CREATIVE CONTENT GENERATION\n"
                "- CASUAL EVERYDAY CONVERSATIONS\n"
                "- USER EDUCATION AND LEARNING SUPPORT"
            ),
            "context": {
                "goal": 
                    "ì‚¬ìš©ìì˜ ì§ˆë¬¸ì— ëŒ€í•´ ê°€ì¥ ì •í™•í•˜ê³  ê´€ë ¨ ìˆëŠ” ì •ë³´ë¥¼ ì œê³µí•˜ëŠ” ê²ƒì„ ëª©í‘œë¡œ í•©ë‹ˆë‹¤. "
                    "ì´ë¥¼ ìœ„í•´, ë‹¤ìŒì„ ì¤€ìˆ˜í•˜ì‹­ì‹œì˜¤:\n"
                    "1. **ì‚¬ìš©ì ì˜ë„ë¥¼ íŒŒì•…**: ì§ˆë¬¸ì˜ ë§¥ë½ê³¼ ë°°ê²½ì„ ì‹ ì¤‘íˆ ë¶„ì„í•˜ì—¬ ì‚¬ìš©ì ìš”êµ¬ë¥¼ ì´í•´í•©ë‹ˆë‹¤.\n"
                    "2. **ì •í™•ì„±ê³¼ ì‹ ë¢°ì„± í™•ë³´**: ì œê³µë˜ëŠ” ì •ë³´ëŠ” ê²€ì¦ ê°€ëŠ¥í•˜ê³ , ë…¼ë¦¬ì ìœ¼ë¡œ íƒ€ë‹¹í•´ì•¼ í•©ë‹ˆë‹¤.\n"
                    "3. **ì‚¬ìš©ì ì¹œí™”ì  ë‹µë³€ ì œê³µ**: ê¸°ìˆ ì ì´ê±°ë‚˜ ë³µì¡í•œ ì§ˆë¬¸ì—ëŠ” ë‹¨ê³„ì ìœ¼ë¡œ ì ‘ê·¼í•˜ë©°, í•„ìš”í•œ ê²½ìš° ê°„ë‹¨í•œ ì–¸ì–´ì™€ ì˜ˆì‹œë¥¼ ì‚¬ìš©í•´ ì„¤ëª…í•©ë‹ˆë‹¤.\n"
                    "4. **ì ì‘ì„± ìœ ì§€**: ì¼ë°˜ ì‚¬ìš©ìì™€ ì „ë¬¸ê°€ ëª¨ë‘ì—ê²Œ ì í•©í•œ ì–¸ì–´ì™€ ì„¸ë¶€ì‚¬í•­ìœ¼ë¡œ ë‹µë³€ì„ ì¡°ì •í•©ë‹ˆë‹¤.\n"
                    "5. **ì°½ì˜ì  ì ‘ê·¼ë²• ì ìš©**: ì‚¬ìš©ìê°€ ìš”ì²­í•˜ëŠ” ë…ì°½ì  ì•„ì´ë””ì–´ë‚˜ ë¬¸ì œ í•´ê²°ì„ ì§€ì›í•©ë‹ˆë‹¤.\n"
                    "6. **í¬ê´„ì  ë¬¸ì œ í•´ê²°**: ê°„ë‹¨í•œ ì§ˆë¬¸ë¿ë§Œ ì•„ë‹ˆë¼ ë³µì¡í•œ ë¬¸ì œë„ ë‹¨ê³„ë³„ë¡œ ì ‘ê·¼í•˜ì—¬ ì‚¬ìš©ìì—ê²Œ ëª…í™•í•˜ê³  ì‹¤í–‰ ê°€ëŠ¥í•œ ë‹µë³€ì„ ì œê³µí•©ë‹ˆë‹¤."
            },
            "conversation": history if isinstance(history, list) else []
        }


        prompt_template["conversation"].append({"role": "user", "content": new_message})
        if relevant_chunks:
            prompt_template["context"]["retrieved_chunks"] = relevant_chunks

        response = bedrock_agent_runtime.retrieve_and_generate(
            input={'text': json.dumps(prompt_template)},
            retrieveAndGenerateConfiguration={
                'knowledgeBaseConfiguration': {
                    'knowledgeBaseId': kbId,
                    'modelArn': modelArn
                },
                'type': 'KNOWLEDGE_BASE'
            }
        )

        output_text = response.get('output', {}).get('text', 'No output generated')
        print("Generated response:", output_text)

        return {
            'statusCode': 200,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({
                "responseCode": "200",
                "responseStatus": "OK",
                "resultData": {
                    "message": output_text
                }
            })
        }

    except Exception as e:
        print("Error occurred:", str(e))
        return {
            'statusCode': 500,
            'headers': {
                'Access-Control-Allow-Origin': '*',
                'Content-Type': 'application/json'
            },
            'body': json.dumps({
                'error': str(e)
            })
        }
